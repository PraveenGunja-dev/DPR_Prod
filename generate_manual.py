from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # --- Header ---
    title = doc.add_heading('Digitalized Daily Progress Report (DPR)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive User & Operational Manual')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    doc.add_paragraph('\n')

    # --- Section: What is DPR? ---
    doc.add_heading('1. What is Digitalized DPR?', level=1)
    doc.add_paragraph(
        "DPR stands for Daily Progress Report. In the context of large-scale projects, it is the primary record of all physical work, manpower utilization, and material usage completed on a project site within a 24-hour period."
    )
    
    doc.add_heading('Why Digitalize?', level=2)
    p = doc.add_paragraph('Traditional manual reporting is prone to delays, data silos, and transcription errors. The Digitalized DPR system solves these issues through:')
    p.style = 'List Bullet'
    doc.add_paragraph('Real-Time Data Entry: Field supervisors enter data directly on the platform.', style='List Bullet')
    doc.add_paragraph('Seamless Validation: Site PMs review entries with cell-level feedback.', style='List Bullet')
    doc.add_paragraph('Direct P6 Integration: Validated data is pushed directly to Oracle P6, eliminating double entry.', style='List Bullet')
    doc.add_paragraph('Centralized Audit Trail: Every change, comment, and approval is tracked.', style='List Bullet')

    # --- Section: User Roles & Workflows ---
    doc.add_heading('2. User Roles & Workflows', level=1)
    
    doc.add_heading('Supervisor (The Field Expert)', level=2)
    doc.add_paragraph(
        "Supervisors are the foundation of the data. They record today's achievements against yesterday's benchmarks."
    )
    doc.add_paragraph('Key Tasks:', style='List Number')
    doc.add_paragraph('Select Project & Sync: Ensure the activity list is up-to-date from the P6 database.', style='List Number')
    doc.add_paragraph('Data Entry: Navigate through various sheet tabs (DP Qty, Manpower, etc.) to record progress.', style='List Number')
    doc.add_paragraph('Submission: Submit completed sheets for PM review.', style='List Number')

    doc.add_heading('Site PM (The Validator)', level=2)
    doc.add_paragraph(
        "Site Project Managers act as the quality gates, ensuring data accuracy before it moves to project controls."
    )
    doc.add_paragraph('Key Tasks:', style='List Number')
    doc.add_paragraph('Dashboard Monitoring: View pending submissions from all site supervisors.', style='List Number')
    doc.add_paragraph('Detailed Review: Use cell-level comments to request specific corrections.', style='List Number')
    doc.add_paragraph('Approval: Confirm accuracy and send the sheet to the PMAG team.', style='List Number')

    doc.add_heading('PMAG (Project Management AG)', level=2)
    doc.add_paragraph(
        "PMAG oversees the entire project data lifecycle and manages the interface with scheduling tools like Oracle P6."
    )
    doc.add_paragraph('Key Tasks:', style='List Number')
    doc.add_paragraph('Final Review: A secondary check to ensure global project standards.', style='List Number')
    doc.add_paragraph('Push to P6: The click of a button that updates the project schedule with real-world data.', style='List Number')

    # --- Section: Understanding the Tables ---
    doc.add_heading('3. Understanding the Reporting Sheets', level=1)
    
    doc.add_heading('DP Qty (Daily Progress Quantity)', level=2)
    doc.add_paragraph(
        "Captures the physical work done (e.g., Cubic Meters of concrete, count of piles). It tracks 'Today Achievement' and automatically calculates 'Cumulative Progress' and 'Balance Scope'."
    )
    
    doc.add_heading('Manpower Details', level=2)
    doc.add_paragraph(
        "Critical for resource monitoring. Tracks direct and indirect labor counts by contractor and site section."
    )

    doc.add_heading('Vendor Block/IDT', level=2)
    doc.add_paragraph(
        "Tracks specific vendor performance and the movement of Internal Delivery Tokens (IDT) for material procurement."
    )

    # --- Section: Common Usage Scenarios ---
    doc.add_heading('4. Common Usage Scenarios', level=1)
    
    doc.add_heading('Scenario: Handling a Rejected Sheet', level=2)
    doc.add_paragraph(
        "If a PM rejects your sheet, it will reappear in your 'Supervisor Dashboard' with a 'Rejected' status. "
        "Open the sheet to see red markers on specific cells. Hover over or click these cells to see the review comments. "
        "Correct the data and resubmit immediately."
    )

    doc.add_heading('Scenario: Editing Past Data', level=2)
    doc.add_paragraph(
        "Data accuracy is paramount. If you need to edit a previous day's report, select the date. "
        "The system will require a 'Business Reason' for the modification, which will be logged for PM review."
    )

    # --- Footer ---
    doc.add_paragraph('\n')
    footer = doc.add_paragraph('Digitalized DPR System - Ensuring Data-Driven Project Excellence.')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    filename = 'Digitalized_DPR_User_Manual.docx'
    doc.save(filename)
    return filename

if __name__ == '__main__':
    try:
        path = create_manual()
        print(f"SUCCESS: Manual generated at {path}")
    except Exception as e:
        print(f"ERROR: {str(e)}")
